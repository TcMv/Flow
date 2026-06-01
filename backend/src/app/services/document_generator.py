"""Document Generator — renders workflow output as PDF and DOCX.

Each renderer produces a branded, professionally formatted document
suitable for government and enterprise use.
"""

from __future__ import annotations

import io
import os
import re
from datetime import datetime, timezone
from io import BytesIO
from typing import Literal

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm, cm
from reportlab.lib.colors import HexColor
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    PageBreak, Frame, PageTemplate, ListFlowable, ListItem,
)
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
from reportlab.platypus.flowables import HRFlowable


# ── Brand Config ─────────────────────────────────────────────────────

BRAND_COLOR = HexColor("#0F6B50")  # Flow emerald
BRAND_COLOR_LIGHT = HexColor("#E8F9F3")
DARK_COLOR = HexColor("#1A1D23")
MUTED_COLOR = HexColor("#5C6068")
BORDER_COLOR = HexColor("#E2E4E8")

from pathlib import Path

# Document storage — ephemeral on Vercel (use /tmp), persistent on-prem
DOCUMENTS_DIR = Path(os.environ.get("FLOW_DOCUMENTS_DIR", "/tmp/flow-documents"))

# ── PDF Generator ────────────────────────────────────────────────────


def _build_pdf_styles():
    """Build a stylesheet for government-branded PDFs."""
    styles = getSampleStyleSheet()

    styles.add(ParagraphStyle(
        "CoverTitle", parent=styles["Title"],
        fontSize=26, leading=32, spaceAfter=6,
        textColor=DARK_COLOR, fontName="Helvetica-Bold",
        alignment=TA_LEFT,
    ))
    styles.add(ParagraphStyle(
        "CoverSubtitle", parent=styles["Normal"],
        fontSize=13, leading=18, spaceAfter=24,
        textColor=MUTED_COLOR, fontName="Helvetica",
        alignment=TA_LEFT,
    ))
    styles.add(ParagraphStyle(
        "SectionHeader", parent=styles["Heading1"],
        fontSize=16, leading=22, spaceBefore=18, spaceAfter=8,
        textColor=BRAND_COLOR, fontName="Helvetica-Bold",
        borderPadding=(0, 0, 4, 0),
    ))
    styles.add(ParagraphStyle(
        "SubSection", parent=styles["Heading2"],
        fontSize=12, leading=16, spaceBefore=12, spaceAfter=4,
        textColor=DARK_COLOR, fontName="Helvetica-Bold",
    ))
    styles.add(ParagraphStyle(
        "Body", parent=styles["Normal"],
        fontSize=10, leading=15, spaceAfter=6,
        textColor=DARK_COLOR, fontName="Helvetica",
        alignment=TA_JUSTIFY,
    ))
    styles.add(ParagraphStyle(
        "FlowBullet", parent=styles["Normal"],
        fontSize=10, leading=15, spaceAfter=3,
        textColor=DARK_COLOR, fontName="Helvetica",
        leftIndent=20, bulletIndent=8,
    ))
    styles.add(ParagraphStyle(
        "Metadata", parent=styles["Normal"],
        fontSize=8, leading=11, spaceAfter=2,
        textColor=MUTED_COLOR, fontName="Helvetica",
        alignment=TA_LEFT,
    ))
    styles.add(ParagraphStyle(
        "Footer", parent=styles["Normal"],
        fontSize=7, leading=9,
        textColor=MUTED_COLOR, fontName="Helvetica",
        alignment=TA_CENTER,
    ))
    return styles


def _add_pdf_header_footer(canvas, doc):
    """Draw header and footer on each PDF page."""
    canvas.saveState()
    # Header line
    canvas.setStrokeColor(BRAND_COLOR)
    canvas.setLineWidth(0.5)
    canvas.line(2*cm, A4[1] - 1.5*cm, A4[0] - 2*cm, A4[1] - 1.5*cm)
    canvas.setFont("Helvetica", 7)
    canvas.setFillColor(MUTED_COLOR)
    canvas.drawString(2*cm, A4[1] - 1.3*cm, "Flow — Self-hosted AI Agent Platform")
    # Footer
    canvas.setStrokeColor(BORDER_COLOR)
    canvas.setLineWidth(0.3)
    canvas.line(2*cm, 1.5*cm, A4[0] - 2*cm, 1.5*cm)
    canvas.setFont("Helvetica", 7)
    canvas.setFillColor(MUTED_COLOR)
    canvas.drawString(2*cm, 1.2*cm, f"Generated {datetime.now(timezone.utc).strftime('%d %B %Y at %H:%M UTC')}")
    canvas.drawRightString(A4[0] - 2*cm, 1.2*cm, f"Page {doc.page}")
    canvas.restoreState()


def render_pdf(
    title: str,
    sections: list[dict],
    author: str = "Flow Agent",
    filename: str = "document.pdf",
) -> bytes:
    """Render a structured document as PDF.

    Args:
        title: Document title (appears on cover)
        sections: List of dicts with keys:
            - heading: Section heading text
            - content: List of paragraphs (str) or bullet items (dict with ``text``)
            - level: "section" or "subsection" (default: "section")
        author: Author name in metadata
        filename: Suggested filename

    Returns:
        PDF content as bytes.
    """
    buffer = BytesIO()
    styles = _build_pdf_styles()

    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        topMargin=2.5*cm,
        bottomMargin=2.5*cm,
        leftMargin=2.5*cm,
        rightMargin=2.5*cm,
        title=title,
        author=author,
    )

    frame = Frame(
        doc.leftMargin, doc.bottomMargin,
        doc.width, doc.height,
        id="normal",
    )
    doc.addPageTemplates([
        PageTemplate(id="main", frames=frame, onPage=_add_pdf_header_footer),
    ])

    elements: list = []

    # ── Cover ─────────────────────────────────────────────────
    elements.append(Spacer(1, 2*cm))
    elements.append(HRFlowable(
        width="100%", thickness=3,
        color=BRAND_COLOR, spaceAfter=12,
    ))
    elements.append(Paragraph(title, styles["CoverTitle"]))
    elements.append(Paragraph(f"Prepared by {author}", styles["CoverSubtitle"]))
    elements.append(Spacer(1, 0.5*cm))
    elements.append(Paragraph(
        f"Generated: {datetime.now(timezone.utc).strftime('%d %B %Y')}",
        styles["Metadata"],
    ))
    elements.append(Spacer(1, 1.5*cm))

    # ── Sections ─────────────────────────────────────────────
    for section in sections:
        heading = section.get("heading", "")
        level = section.get("level", "section")
        content = section.get("content", [])

        if heading:
            if level == "section":
                elements.append(Paragraph(heading, styles["SectionHeader"]))
            else:
                elements.append(Paragraph(heading, styles["SubSection"]))

        for item in content:
            if isinstance(item, dict) and "text" in item:
                # Bullet item
                elements.append(Paragraph(
                    f"• {item['text']}", styles["FlowBullet"],
                ))
            elif isinstance(item, str):
                if item.startswith("---"):
                    elements.append(HRFlowable(
                        width="100%", thickness=0.5,
                        color=BORDER_COLOR, spaceAfter=8, spaceBefore=8,
                    ))
                elif item == "":
                    elements.append(Spacer(1, 6))
                else:
                    elements.append(Paragraph(item, styles["Body"]))
            elif isinstance(item, list):
                # Table data
                if item and isinstance(item[0], list):
                    table_data = item
                    tbl = Table(table_data, repeatRows=1)
                    tbl.setStyle(TableStyle([
                        ("BACKGROUND", (0, 0), (-1, 0), BRAND_COLOR),
                        ("TEXTCOLOR", (0, 0), (-1, 0), HexColor("#FFFFFF")),
                        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                        ("FONTSIZE", (0, 0), (-1, -1), 9),
                        ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("GRID", (0, 0), (-1, -1), 0.5, BORDER_COLOR),
                        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [HexColor("#FFFFFF"), HexColor("#F8FAFC")]),
                        ("TOPPADDING", (0, 0), (-1, -1), 4),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                        ("LEFTPADDING", (0, 0), (-1, -1), 6),
                    ]))
                    elements.append(tbl)
                    elements.append(Spacer(1, 8))

    doc.build(elements)
    return buffer.getvalue()


# ── DOCX Generator ───────────────────────────────────────────────────


def render_docx(
    title: str,
    sections: list[dict],
    author: str = "Flow Agent",
) -> bytes:
    """Render a structured document as DOCX.

    Args:
        title: Document title
        sections: Same structure as ``render_pdf``
        author: Author name

    Returns:
        DOCX content as bytes.
    """
    doc = Document()

    # ── Styles ───────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # ── Cover ─────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x0F, 0x6B, 0x50)

    p = doc.add_paragraph()
    run = p.add_run(f"Prepared by {author}")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x5C, 0x60, 0x68)

    p = doc.add_paragraph()
    run = p.add_run(f"Generated: {datetime.now(timezone.utc).strftime('%d %B %Y')}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x78, 0x7C, 0x85)

    doc.add_paragraph()  # spacer

    # ── Sections ─────────────────────────────────────────────
    for section in sections:
        heading = section.get("heading", "")
        level = section.get("level", "section")
        content = section.get("content", [])

        if heading:
            if level == "section":
                h = doc.add_heading(heading, level=1)
                for run in h.runs:
                    run.font.color.rgb = RGBColor(0x0F, 0x6B, 0x50)
            else:
                h = doc.add_heading(heading, level=2)
                for run in h.runs:
                    run.font.color.rgb = RGBColor(0x1A, 0x1D, 0x23)

        for item in content:
            if isinstance(item, dict) and "text" in item:
                p = doc.add_paragraph(item["text"], style="List Bullet")
            elif isinstance(item, str):
                if item == "":
                    doc.add_paragraph()
                elif item.startswith("---"):
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after = Pt(6)
                    run = p.add_run("─" * 60)
                    run.font.size = Pt(6)
                    run.font.color.rgb = RGBColor(0xE2, 0xE4, 0xE8)
                else:
                    doc.add_paragraph(item)
            elif isinstance(item, list) and item and isinstance(item[0], list):
                # Table
                table = doc.add_table(rows=len(item), cols=len(item[0]))
                table.style = "Light Grid Accent 1"
                for i, row_data in enumerate(item):
                    for j, cell_text in enumerate(row_data):
                        cell = table.rows[i].cells[j]
                        cell.text = str(cell_text)
                        if i == 0:
                            for p in cell.paragraphs:
                                for r in p.runs:
                                    r.bold = True
                                    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ── Public API ────────────────────────────────────────────────────────


def render_document(
    title: str,
    sections: list[dict],
    format: Literal["pdf", "docx"] = "pdf",
    author: str = "Flow Agent",
) -> tuple[bytes, str, str]:
    """Render a structured document and return (content, filename, mime_type).

    Args:
        title: Document title
        sections: Structured sections (see ``render_pdf``)
        format: Output format — "pdf" or "docx"
        author: Author name

    Returns:
        Tuple of (bytes content, filename with extension, MIME type).
    """
    safe_title = re.sub(r"[^\w\s-]", "", title).strip().replace(" ", "-")[:50]

    if format == "docx":
        content = render_docx(title, sections, author)
        filename = f"{safe_title}.docx"
        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    else:
        content = render_pdf(title, sections, author, filename=f"{safe_title}.pdf")
        filename = f"{safe_title}.pdf"
        mime = "application/pdf"

    return content, filename, mime


# ── Workflow Output → Document Sections ──────────────────────────────


def workflow_output_to_sections(
    title: str,
    task_results: dict,
    output_defs: list[str] | None = None,
) -> list[dict]:
    """Convert workflow task results into structured document sections.

    Args:
        title: Document title
        task_results: Dict of ``{task_id: output_data}`` from a workflow run
        output_defs: Optional list of output descriptions from the workflow def

    Returns:
        List of section dicts suitable for ``render_document``.
    """
    sections = []

    # Cover / intro
    sections.append({
        "heading": "Executive Summary",
        "level": "section",
        "content": [
            f"This document was generated by Flow from the workflow \"{title}\".",
            f"Completed at {datetime.now(timezone.utc).strftime('%d %B %Y at %H:%M UTC')}.",
            "",
        ],
    })

    if output_defs:
        sections.append({
            "heading": "Outputs Produced",
            "level": "section",
            "content": [{"text": d} for d in output_defs] + [""],
        })

    # Task results
    sections.append({
        "heading": "Task Results",
        "level": "section",
        "content": [],
    })

    for task_id, output in sorted(task_results.items()):
        if not output:
            continue

        summary = output.get("summary", "")
        result_text = output.get("result", "")
        status = output.get("status", "completed")

        if status == "failed":
            content = [f"⚠️ This task did not complete successfully."]
            if output.get("error"):
                content.append(output["error"])
        else:
            content = []
            if summary:
                content.append(summary)
            if result_text and result_text != summary:
                content.append(result_text)

        if content:
            sections.append({
                "heading": task_id.replace("_", " ").title(),
                "level": "subsection",
                "content": content,
            })

    return sections
